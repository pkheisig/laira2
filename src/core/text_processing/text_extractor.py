"""
Text Extractor Module for LAIRA.

This module provides the TextExtractor class for extracting text from various document formats
including PDF, DOCX, and plain text files.
"""

import os
import logging
from typing import List, Dict, Any, Optional, Union, BinaryIO, TextIO
import io
import re # Added for section splitting in PDF
import time

# Import document processing libraries
import pypdf
# pdfplumber is removed as fitz handles text extraction
import docx
import fitz  # PyMuPDF
from PIL import Image as PILImage
from langchain.schema import Document

# Import Vertex AI related libraries
from google.cloud import aiplatform
import vertexai
from vertexai.generative_models import GenerativeModel, Part, Image as VertexImage, GenerationConfig

# Import utility functions
from .utils import detect_file_type

# Set up logging
logger = logging.getLogger(__name__)


class TextExtractor:
    """
    A class for extracting text from various document formats.
    
    This class provides methods to extract text from PDF, DOCX, and plain text files.
    It handles file access and parsing, with appropriate error handling.
    It can also analyze images within PDF files using Gemini Vision.
    """
    
    def __init__(self, extraction_config: Optional[Dict[str, Any]] = None):
        """
        Initialize the TextExtractor with optional configuration.
        
        Args:
            extraction_config: Optional configuration dictionary for extraction settings
        """
        self.config = extraction_config or {}
        self.last_error = None
        self.vision_model_name = self.config.get("vision_model_name", "gemini-pro-vision") # Default to gemini-pro-vision
        self.vision_model = None # Initialize as None

        # Lazy load the vision model only when needed (specifically for PDFs)
        # This avoids initializing it if only text/docx files are processed.
        # Ensure Vertex AI is initialized beforehand where this class is used.

    def _initialize_vision_model(self):
        """Initializes the Gemini Vision model if not already done."""
        if self.vision_model is None:
            try:
                # Ensure Vertex AI is initialized beforehand (typically in the main app or processor)
                # aiplatform.init(...) should be called elsewhere
                self.vision_model = GenerativeModel(self.vision_model_name)
                logger.info(f"Gemini Vision model '{self.vision_model_name}' initialized.")
            except Exception as e:
                logger.error(f"Failed to initialize Gemini Vision model '{self.vision_model_name}': {e}")
                self.last_error = f"Vision model init failed: {e}"
                # Keep self.vision_model as None

    def extract_text(self, file_path: str) -> Union[str, List[Document], None]:
        """
        Extract text from a file based on its detected type.
        For PDFs, it returns a list of Document objects (one per page), including image analysis.
        For other types, it returns a single string.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted content as a list of Document objects (for PDF) or a string (for others), 
            or None if extraction failed.
        """
        try:
            # Reset last error
            self.last_error = None
            
            # Check if file exists
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Detect file type
            file_type, _ = detect_file_type(file_path)
            
            # Extract text based on file type
            if file_type == 'pdf':
                # Initialize vision model only if processing a PDF
                self._initialize_vision_model()
                if self.vision_model is None:
                    logger.error("Vision model failed to initialize, cannot process PDF images.")
                    # Optionally, proceed with text-only extraction or raise an error
                    # For now, we will try text extraction but log the warning
                    # return None # Or raise an error
                return self.extract_from_pdf(file_path)
            elif file_type == 'docx':
                return self.extract_from_docx(file_path)
            elif file_type == 'txt':
                return self.extract_from_txt(file_path)
            else:
                error_msg = f"Unsupported file type: {file_type}"
                self.last_error = error_msg
                logger.error(error_msg)
                return None
                
        except Exception as e:
            self.last_error = str(e)
            logger.error(f"Error extracting text from {file_path}: {e}", exc_info=True)
            return None
    
    def extract_from_pdf(self, file_path: str) -> List[Document]:
        """
        Extract text and analyze images from a PDF file using fitz and Gemini Vision.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            A list of Document objects, one per page, containing text and image analysis.
            
        Raises:
            FileNotFoundError: If the file does not exist
            IOError: If there is an error reading or processing the file
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"PDF file not found: {file_path}")
        
        documents = []
        try:
            # Open PDF with PyMuPDF
            with fitz.open(file_path) as pdf:
                total_pages = len(pdf)
                logger.info(f"Processing PDF: {os.path.basename(file_path)} ({total_pages} pages)")
                for page_num, page in enumerate(pdf, 1):
                    page_start_time = time.time()
                    try:
                        # Extract text
                        logger.debug(f"Extracting text from page {page_num}...")
                        text = page.get_text("text") or "" # Ensure text is not None, use "text" for better structure

                        # --- Image Extraction and Analysis ---
                        image_analysis_results = []
                        image_list = []
                        try:
                            image_list = page.get_images(full=True)
                            logger.debug(f"Found {len(image_list)} potential images on page {page_num}.")
                        except Exception as img_list_err:
                             logger.warning(f"Could not list images on page {page_num} of {file_path}: {img_list_err}")

                        if self.vision_model and image_list: # Only process if model initialized and images exist
                            logger.info(f"Analyzing {len(image_list)} images on page {page_num}...")
                            for img_idx, img_info in enumerate(image_list):
                                img_start_time = time.time()
                                try:
                                    xref = img_info[0]  # Image reference
                                    # Check if image is valid (e.g., has non-zero dimensions)
                                    if img_info[2] <= 0 or img_info[3] <= 0:
                                        logger.warning(f"Skipping invalid image {img_idx + 1} (xref: {xref}) on page {page_num} due to zero dimensions.")
                                        continue

                                    logger.debug(f"Extracting image {img_idx + 1} (xref: {xref}) on page {page_num}...")
                                    base_image = pdf.extract_image(xref)

                                    if base_image:
                                        image_bytes = base_image["image"]
                                        logger.debug(f"Image {img_idx + 1} size: {len(image_bytes)} bytes, format: {base_image.get('ext', 'unknown')}")

                                        # Convert to PIL Image for processing
                                        image = PILImage.open(io.BytesIO(image_bytes))

                                        # Convert CMYK/P/RGBA to RGB (essential for Gemini)
                                        if image.mode != 'RGB':
                                            logger.debug(f"Converting image {img_idx + 1} from mode {image.mode} to RGB.")
                                            try:
                                                image = image.convert('RGB')
                                            except Exception as convert_err:
                                                 logger.warning(f"Failed to convert image {img_idx+1} to RGB: {convert_err}. Skipping analysis.")
                                                 continue

                                        # Save to bytes for Gemini (using JPEG is generally good)
                                        img_byte_arr = io.BytesIO()
                                        image.save(img_byte_arr, format='JPEG', quality=90) # Use JPEG for compatibility
                                        img_bytes_for_gemini = img_byte_arr.getvalue()
                                        logger.debug(f"Prepared image {img_idx + 1} for Gemini analysis ({len(img_bytes_for_gemini)} bytes).")

                                        # Create Gemini image part
                                        vertex_image = VertexImage.from_bytes(img_bytes_for_gemini)

                                        # Analyze with Gemini
                                        prompt = """Analyze this scientific figure/image in detail. Focus on:
1. Type of visualization (graph, diagram, microscopy image, flow chart, etc.)
2. Key elements, labels, and data shown.
3. Main findings or patterns visible in the image.
4. Relationship to scientific concepts if evident.
5. Clear and concise description of the figure's scientific conclusion.

Provide a concise but thorough description suitable for understanding the figure's purpose and content."""

                                        # Configure generation - potentially make these configurable
                                        gen_config = GenerationConfig(
                                            temperature=0.2, # Low temp for factual description
                                            max_output_tokens=500, # Limit description length
                                        )

                                        logger.info(f"Sending image {img_idx + 1} (page {page_num}) to Gemini Vision...")
                                        response = self.vision_model.generate_content(
                                            [prompt, vertex_image],
                                            generation_config=gen_config,
                                            # stream=False # Ensure non-streaming for this use case
                                        )

                                        # Handle potential safety blocks or empty responses
                                        analysis_text = "Analysis could not be generated." # Default
                                        if response.candidates:
                                             candidate = response.candidates[0]
                                             if candidate.content and candidate.content.parts:
                                                  analysis_text = candidate.content.parts[0].text or analysis_text
                                             # Log safety ratings if needed
                                             # logger.debug(f"Image {img_idx+1} safety ratings: {candidate.safety_ratings}")

                                        logger.info(f"Received analysis for image {img_idx + 1} (page {page_num}).")
                                        image_analysis_results.append(f"[FIGURE {img_idx + 1} ANALYSIS]: {analysis_text}")
                                        img_duration = time.time() - img_start_time
                                        logger.debug(f"Image {img_idx + 1} analysis took {img_duration:.2f}s.")

                                    else:
                                        logger.warning(f"Could not extract image data for xref {xref} on page {page_num}.")

                                except Exception as img_err:
                                    logger.error(f"Error processing image {img_idx + 1} (xref: {xref}) on page {page_num} of {file_path}: {img_err}", exc_info=True)
                                    continue # Skip this image if analysis fails
                        else:
                             logger.debug(f"No images found or vision model not available for page {page_num}.")


                        # --- Text Processing (Basic Section Detection) ---
                        logger.debug("Processing text content for sections...")
                        lines = text.split('\\n')
                        processed_lines = []
                        current_section = "Unknown" # Default section

                        for line in lines:
                            stripped_line = line.strip()
                            # Simple heuristic for headers: all caps, few words, not too long
                            if stripped_line.isupper() and 0 < len(stripped_line.split()) <= 7 and len(stripped_line) < 80:
                                current_section = stripped_line
                                processed_lines.append(f"\\n### {current_section} ###\\n")
                                logger.debug(f"Detected section header: {current_section}")
                            else:
                                # Clean excess whitespace but preserve line breaks that might indicate paragraphs
                                cleaned_line = ' '.join(stripped_line.split())
                                if cleaned_line:
                                    processed_lines.append(cleaned_line)

                        # Join text content with paragraph spacing
                        processed_text = '\\n\\n'.join(processed_lines)

                        # Append image analyses to the text if any
                        if image_analysis_results:
                            logger.debug(f"Appending {len(image_analysis_results)} image analyses to page {page_num} text.")
                            image_descriptions = "\\n\\n### FIGURES AND VISUALIZATIONS ###\\n" + "\\n\\n".join(image_analysis_results)
                            processed_text += "\\n\\n" + image_descriptions

                        # Create base metadata
                        base_meta = {
                            'source': file_path,
                            'filename': os.path.basename(file_path),
                            'page': page_num,
                            'total_pages': total_pages,
                            'section': current_section,
                            'has_images': bool(image_list),
                            'image_count': len(image_list),
                            'analyzed_image_count': len(image_analysis_results)
                        }
                        # 1) Text-only chunk
                        text_meta = base_meta.copy()
                        text_meta['chunk_type'] = 'text'
                        doc_text = Document(page_content=processed_text, metadata=text_meta)
                        documents.append(doc_text)
                        # 2) Separate figure analysis chunks
                        for idx, analysis in enumerate(image_analysis_results, 1):
                            fig_meta = base_meta.copy()
                            fig_meta['chunk_type'] = 'figure'
                            fig_meta['figure_index'] = idx
                            doc_fig = Document(page_content=analysis, metadata=fig_meta)
                            documents.append(doc_fig)
                        page_duration = time.time() - page_start_time
                        logger.debug(f"Page {page_num} processed in {page_duration:.2f}s.")

                    except Exception as page_err:
                         logger.error(f"Error processing page {page_num} of {file_path}: {page_err}", exc_info=True)
                         # Create a placeholder document indicating the error for this page
                         documents.append(Document(
                              page_content=f"[ERROR PROCESSING PAGE {page_num}: {page_err}]",
                              metadata={'source': file_path, 'page': page_num, 'error': str(page_err)}
                         ))
                         continue # Skip to the next page

        except (fitz.errors.FileNotFoundError, FileNotFoundError):
             logger.error(f"PDF file not found: {file_path}")
             raise FileNotFoundError(f"PDF file not found: {file_path}")
        except Exception as e:
            # Catch errors opening the file or during iteration setup
            logger.error(f"Failed to process PDF {file_path}: {str(e)}", exc_info=True)
            raise IOError(f"Failed to process PDF {file_path}: {str(e)}")

        logger.info(f"Extracted {len(documents)} page(s) from {os.path.basename(file_path)}.")
        return documents
    
    def extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text as a string
            
        Raises:
            FileNotFoundError: If the file does not exist
            ValueError: If the file is not a valid DOCX
            IOError: If there is an error reading the file
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    full_text.append(" | ".join(row_text))
            
            return "\n".join(full_text)
        except Exception as e:
            raise IOError(f"Failed to extract text from DOCX: {e}")
    
    def extract_from_txt(self, file_path: str) -> str:
        """
        Extract text from a plain text file.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Extracted text as a string
            
        Raises:
            FileNotFoundError: If the file does not exist
            IOError: If there is an error reading the file
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Text file not found: {file_path}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encodings if UTF-8 fails
            encodings = ['latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, raise an error
            raise IOError(f"Failed to decode text file with any known encoding")
        except Exception as e:
            raise IOError(f"Failed to read text file: {e}")
    
    def get_document_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from a document.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary containing metadata
        """
        metadata = {
            'filename': os.path.basename(file_path),
            'file_size': os.path.getsize(file_path),
            'file_type': None,
            'mime_type': None,
            'page_count': None,
            'creation_date': None,
            'modification_date': None,
        }
        
        try:
            # Get file type and mime type
            file_type, mime_type = detect_file_type(file_path)
            metadata['file_type'] = file_type
            metadata['mime_type'] = mime_type
            
            # Get file dates
            metadata['creation_date'] = os.path.getctime(file_path)
            metadata['modification_date'] = os.path.getmtime(file_path)
            
            # Get page count for PDF
            if file_type == 'pdf':
                try:
                    with open(file_path, 'rb') as file:
                        reader = pypdf.PdfReader(file)
                        metadata['page_count'] = len(reader.pages)
                except Exception as e:
                    logger.warning(f"Failed to get PDF page count: {e}")
            
            # Get page count for DOCX
            elif file_type == 'docx':
                try:
                    doc = docx.Document(file_path)
                    metadata['page_count'] = len(doc.sections)
                except Exception as e:
                    logger.warning(f"Failed to get DOCX page count: {e}")
            
        except Exception as e:
            logger.error(f"Error getting document metadata: {e}")
        
        return metadata